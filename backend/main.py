# import math
import psutil
from fastapi import FastAPI, Header, Depends, HTTPException, Response, Request, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
from transformers import pipeline
from passlib.context import CryptContext
from jose import jwt
import database 
import docx
import io
from docx import Document
import fitz  # PyMuPDF
import bcrypt
# from sentence_transformers import SentenceTransformer, util
import re
import tempfile
import os
# import json
# from transformers import pipeline, AutoModelForCausalLM, AutoTokenizer
import csv
# import ollama
import torch
from striprtf.striprtf import rtf_to_text
import re
import subprocess
import platform
import pandas as pd

IS_WINDOWS = platform.system() == "Windows"
if IS_WINDOWS:
    try:
        import pythoncom
        import win32com.client
        from docx2pdf import convert
    except ImportError:
        pass


os.environ['TRANSFORMERS_CACHE'] = '/tmp/huggingface_cache'
os.environ['SENTENCE_TRANSFORMERS_HOME'] = '/tmp/sentence_transformers_cache'
torch.set_num_threads(1)
torch.set_grad_enabled(False)

database.Base.metadata.create_all(bind=database.engine)

app = FastAPI()

# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["http://localhost:5173"], 
#     allow_credentials=True, # Это разрешает передачу кук
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

IS_PRODUCTION = os.getenv("IS_PRODUCTION", "false").lower() == "true"
frontend_url = os.getenv("FRONTEND_URL", "https://ai-sentinel-ppd9-chi.vercel.app")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        frontend_url,
        "https://ai-sentinel-ppd9-chi.vercel.app",
        "http://localhost:5173"
    ], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Загружаем модель для сравнения смыслов (около 400 МБ)
# print("Загрузка семантической модели...")
# semantic_model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')

# model_path = os.path.abspath("./studying/ai_detector_model4")
# if os.path.exists(model_path):
#     ai_classifier = pipeline(
#         "text-classification",
#         model=model_path,
#         tokenizer=model_path,
#         device=0,
#         truncation=True,
#         max_length=512,
#     )
# else:
#     ai_classifier = None

# Настройки безопасности
SECRET_KEY = "SUPER_SECRET_KEY_FOR_DIPLOMA"
ALGORITHM = "HS256"
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# Хелперы
def get_password_hash(password: str):
    # Превращаем пароль в байты, генерируем соль и хешируем
    pwd_bytes = password.encode('utf-8')
    salt = bcrypt.gensalt()
    hashed_password = bcrypt.hashpw(pwd_bytes, salt)
    return hashed_password.decode('utf-8') # Сохраняем как строку

def verify_password(plain_password: str, hashed_password: str):
    password_byte_enc = plain_password.encode('utf-8')
    hashed_password_enc = hashed_password.encode('utf-8')
    return bcrypt.checkpw(password_byte_enc, hashed_password_enc)

def get_db():
    db = database.SessionLocal()
    try:
        yield db
    finally:
        db.close()

def log_memory(step_name: str):
    mem = psutil.virtual_memory()
    used_mb = mem.used / 1024 / 1024
    total_mb = mem.total / 1024 / 1024
    percent = mem.percent
    # flush=True заставляет Питон выводить текст мгновенно, не дожидаясь буфера!
    print(f"📊 [ПАМЯТЬ | {step_name}] Занято: {used_mb:.0f} MB из {total_mb:.0f} MB ({percent}%)", flush=True)

# Функция получения текущего пользователя из куки
# def get_current_user(request: Request, db: Session = Depends(get_db)):
#     token = request.cookies.get("access_token")
#     if not token:
#         return None
#     try:
#         payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
#         email = payload.get("sub")
#         return db.query(database.User).filter(database.User.email == email).first()
#     except:
#         return None
def get_current_user(authorization: str = Header(None), db: Session = Depends(get_db)):
    if not authorization or not authorization.startswith("Bearer "):
        return None
    token = authorization.split(" ")[1] # Берем токен после "Bearer "
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email = payload.get("sub")
        return db.query(database.User).filter(database.User.email == email).first()
    except:
        return None
    
def check_formatting(file_bytes, settings: database.CheckSettings):
    errors = []
    try:
        doc = Document(io.BytesIO(file_bytes))
        
        # --- 1. ПРОВЕРКА ПОЛЕЙ СТРАНИЦЫ [Поля] ---
        if doc.sections:
            section = doc.sections[0]
            # Округляем до 1 знака (библиотека выдает в см через .cm)
            m_left = round(section.left_margin.cm, 1)
            m_right = round(section.right_margin.cm, 1)
            m_top = round(section.top_margin.cm, 1)
            m_bottom = round(section.bottom_margin.cm, 1)

            if abs(m_left - settings.margin_left) > 0.1:
                errors.append(f"[Поля] Левое поле: {m_left} см (нужно {settings.margin_left})")
            if abs(m_right - settings.margin_right) > 0.1:
                errors.append(f"[Поля] Правое поле: {m_right} см (нужно {settings.margin_right})")
            if abs(m_top - settings.margin_top) > 0.1:
                errors.append(f"[Поля] Верхнее поле: {m_top} см (нужно {settings.margin_top})")
            if abs(m_bottom - settings.margin_bottom) > 0.1:
                errors.append(f"[Поля] Нижнее поле: {m_bottom} см (нужно {settings.margin_bottom})")

        # --- 2. ПРОВЕРКА ТИПОГРАФИКИ [Типографика] ---
        font_err_shown = False
        size_err_shown = False
        
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text or len(text) < 20: continue # Пропускаем заголовки и пустые строки

            # Проверяем каждый кусок текста (run) в абзаце
            for run in p.runs:
                if not run.text.strip(): continue
                
                # Имя шрифта (берем из run или из стиля абзаца)
                f_name = run.font.name or p.style.font.name
                # Размер шрифта
                f_size = run.font.size.pt if run.font.size else (p.style.font.size.pt if p.style.font.size else None)

                if f_name and f_name != settings.font_name and not font_err_shown:
                    errors.append(f"[Типографика] Шрифт: обнаружен {f_name} (нужен {settings.font_name})")
                    font_err_shown = True
                
                if f_size and round(f_size) != settings.font_size and not size_err_shown:
                    errors.append(f"[Типографика] Размер текста: {round(f_size)}pt (нужен {settings.font_size}pt)")
                    size_err_shown = True
            
            if font_err_shown and size_err_shown: break

        # --- 3. ПРОВЕРКА БИБЛИОГРАФИИ [Библиография] ---
        refs_count = 0
        found_ref_header = False
        for p in doc.paragraphs:
            t = p.text.lower().strip()
            # Маркер начала списка литературы
            if any(x in t for x in ["список литературы", "библиографическ", "references"]):
                found_ref_header = True
                continue
            
            # Считаем наполненные строки после заголовка
            if found_ref_header and len(t) > 10:
                # Если строка начинается с цифры или скобки [1] - это точно ссылка
                if re.match(r'^[\(\[]?\d+[\)\]\.]', t) or len(t) > 30:
                    refs_count += 1
        
        if refs_count < settings.min_references:
            errors.append(f"[Библиография] Источники: найдено {refs_count} (минимум {settings.min_references})")

    except Exception as e:
        print(f"Error in check_formatting: {e}")
        errors.append("[Система] Ошибка при чтении структуры DOCX")
    
    return errors

# gpt2_tokenizer = None
# gpt2_model = None

# def get_gpt2_models():
#     """Ленивая загрузка GPT-2 только при необходимости"""
#     global gpt2_tokenizer, gpt2_model
#     if gpt2_tokenizer is None or gpt2_model is None:
#         print("Эконом-загрузка GPT-2 для перплексии...", flush=True)
#         model_name = "gpt2"
#         gpt2_tokenizer = AutoTokenizer.from_pretrained(model_name)
#         # low_cpu_mem_usage=True помогает не раздувать память при загрузке
#         gpt2_model = AutoModelForCausalLM.from_pretrained(model_name, low_cpu_mem_usage=True)
#     return gpt2_tokenizer, gpt2_model

# def calculate_perplexity(text):
#     tokenizer, model = get_gpt2_models()
#     encodings = tokenizer(
#         text,
#         return_tensors="pt",
#         truncation=True,
#         max_length=1024
#     )

#     input_ids = encodings.input_ids

#     with torch.no_grad():
#         outputs = model(input_ids, labels=input_ids)

#     loss = outputs.loss.item()   # tensor → float
#     perplexity = math.exp(loss)  # float

#     return perplexity

def split_into_chunks(text, chunk_size=4000):

    words = text.split()
    chunks = []

    for i in range(0, len(words), chunk_size):
        chunk = " ".join(words[i:i+chunk_size])
        chunks.append(chunk)

    return chunks

# def run_ai_logic(text: str):
#     # Очистка и ограничение текста (для скорости берем первые ~4000 символов)
#     # Лама понимает и больше, но для быстрой детекции этого хватит
#     sample_text = " ".join(text.split())[:4000]

#     prompt = f"""
#     Ты — эксперт-криминалист по цифровым текстам. Проанализируй фрагмент научной работы.
#     Определи: сгенерирован ли он ИИ (ChatGPT, Llama и др.) или написан человеком.
    
#     Обрати внимание на:
#     1. Избыточную структурированность и "идеальность" списков.
#     2. Типичные вводные фразы ИИ ("важно отметить", "таким образом", "в заключение стоит сказать").
#     3. Однообразие синтаксических конструкций.

#     Текст для анализа:
#     {sample_text}

#     Верни ответ СТРОГО в формате JSON:
#     {{
#       "label": "AI" или "Human",
#       "score": вероятность ИИ от 0.0 до 1.0,
#     }}
#     """

#     try:
#         # Вызов локальной модели
#         response = ollama.chat(
#             model='llama3.1:8b', 
#             messages=[{'role': 'user', 'content': prompt}],
#             format='json' # Заставляем модель выдать чистый JSON
#         )
        
#         # Парсим результат
#         result = json.loads(response['message']['content'])
        
#         # Приводим к формату, который ждет твой фронтенд и база
#         label = result.get("label", "Human")
#         score = float(result.get("score", 0.0))

#         return label, score, 1

#     except Exception as e:
#         print(f"Ollama Error: {e}")
#         return "Error", 0.0, 0, f"Ошибка локальной нейросети: {str(e)}"

# def run_ai_logic(text: str):
#     # Берем текст (не более 512 токенов, так как DistilBERT плохо работает с 15000)
#     sample_text = " ".join(text.split())[:512]

#     # 1. Получаем предсказание от DistilBERT
#     # Модель выдаст список: [{'label': 'LABEL_0', 'score': 0.95}]
#     result = ai_classifier(sample_text)[0]
    
#     # Предположим: LABEL_1 - это AI, LABEL_0 - это Human
#     # Если ваша модель выдала LABEL_1, значит это AI.
#     # Если результат LABEL_0, score 0.9, значит вероятность AI = 0.1
    
#     # Приводим к единой шкале (вероятность того, что это ИИ)
#     if result['label'] == 'LABEL_1':
#         ai_prob = result['score']
#     else:
#         ai_prob = 1.0 - result['score']

#     # 2. Perplexity (оставляем как было)
#     ppl = calculate_perplexity(sample_text)
#     if ppl < 25:
#         ppl_score = 0.8
#     elif ppl < 40:
#         ppl_score = 0.5
#     else:
#         ppl_score = 0.2

#     # 3. Комбинирование (вес BERT теперь важнее — 0.7)
#     final_score = (ai_prob * 0.7) + (ppl_score * 0.3)

#     label = "AI" if final_score > 0.5 else "Human"

#     return label, final_score, 1

def clean_text_thoroughly(text: str) -> str:
    # 1. Заменяем переносы строк на пробелы
    text = text.replace('\n', ' ').replace('\r', ' ')
    # 2. split() без аргументов делит строку по ЛЮБОМУ количеству пробельных символов
    # 3. join() собирает слова обратно через ОДИН пробел
    return " ".join(text.split())

def run_ai_logic(text: str):
    import os
    import requests
    import json
    
    print("🚀 Запуск детектора ИИ через Облачный API (0 MB RAM)...", flush=True)
    
    try:
        # 1. Очистка и разбивка текста
        clean_text = clean_text_thoroughly(text)
        words = clean_text.split()
        
        if not words:
            return "Human", 0.0, 0
            
        chunk_size = 300
        overlap = 50
        chunks = [" ".join(words[i:i+chunk_size]) for i in range(0, len(words), chunk_size - overlap)]
        
        # Для API берем максимум 10 равномерных чанков, чтобы не спамить сервер
        if len(chunks) > 10:
            step = len(chunks) / 10
            chunks = [chunks[int(i * step)] for i in range(10)]

        # 2. Обращение к API Hugging Face
        API_URL = "https://api-inference.huggingface.co/models/ss1ree/ai-sentinel-model"
        
        hf_token = os.getenv("HF_TOKEN", "")
        # Если токена нет, Hugging Face часто отдает 403 или 404
        headers = {"Authorization": f"Bearer {hf_token}"} 
        
        # Иногда нужно явно указать Content-Type
        headers["Content-Type"] = "application/json"
        
        payload = {
            "inputs": chunks,
            "options": {"wait_for_model": True}
        }
        
        # Используем POST, как и было, но добавим проверку токена
        response = requests.post(API_URL, headers=headers, json=payload, timeout=60)
        if response.status_code != 200:
            print(f"DEBUG: Status {response.status_code}, Response: {response.text}", flush=True)
            return "Error", 0.0, 0
        
        if response.status_code != 200:
            print(f"Ошибка API Hugging Face: {response.text}", flush=True)
            return "Error", 0.0, 0
            
        results = response.json()
        
        # 3. Обработка ответов
        ai_scores =[]
        for res_list in results:
            # API возвращает список вероятностей для каждого куска
            if isinstance(res_list, list) and len(res_list) > 0:
                best_res = res_list[0]
                lbl = str(best_res.get('label', '')).upper()
                is_ai = lbl in ['LABEL_1', 'AI', '1', 'FAKE']
                score = best_res.get('score', 0.0) if is_ai else (1.0 - best_res.get('score', 0.0))
                ai_scores.append(score)

        if not ai_scores:
            return "Human", 0.0, 0

        # 4. Финальная оценка (ваша фирменная логика)
        ai_chunks_count = sum(1 for s in ai_scores if s > 0.8)
        ai_ratio = ai_chunks_count / len(ai_scores)
        
        k = max(1, int(len(ai_scores) * 0.3))
        top_scores = sorted(ai_scores, reverse=True)[:k]
        base_score = sum(top_scores) / len(top_scores)
        
        final_score = max(base_score, 0.75) if ai_ratio > 0.65 else base_score
        label = "AI" if final_score > 0.65 else "Human"
        
        print("✅ Успешно проверено через API!", flush=True)
        return label, round(float(final_score), 2), len(chunks)

    except Exception as e:
        print(f"Критическая ошибка детектора ИИ: {e}", flush=True)
        return "Error", 0.0, 0

# Функция для извлечения текста
async def extract_text_from_file_bytes(file_bytes: bytes, filename: str):
    content = ""
    extension = filename.split(".")[-1].lower()
    
    try:
        if extension == "docx":
            # Читаем docx из байтов в памяти
            doc = docx.Document(io.BytesIO(file_bytes))
            content = "\n".join([p.text for p in doc.paragraphs])
            
        elif extension == "pdf":
            # Читаем pdf из байтов в памяти
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            for page in doc:
                content += page.get_text()
                
        elif extension == "rtf":
            # Читаем RTF. Используем cp1251 или utf-8 с игнорированием битых байтов
            raw_text = file_bytes.decode("cp1251", errors="ignore")
            content = rtf_to_text(raw_text)
                
        elif extension == "txt":
            content = file_bytes.decode("utf-8")
            
    except Exception as e:
        print(f"Ошибка при извлечении текста: {e}")
        
    return content

def check_semantic_rules(doc, settings: database.CheckSettings):
    import gc
    from sentence_transformers import SentenceTransformer, util
    import re
    
    errors = []
    paragraphs =[p.text.strip() for p in doc.paragraphs if p.text.strip()]
    full_text_lower = "\n".join(paragraphs).lower()
    org_ru, org_en, abstract_text, main_body = "", "", "", ""
    
    for i, p in enumerate(paragraphs):
        p_low = p.lower()
        if i < 15:
            if any(x in p_low for x in["университет", "институт", "academy", "university", "reshetnev"]):
                if re.search('[а-яА-Я]', p): org_ru = p
                else: org_en = p
        if not abstract_text and any(x in p_low for x in["аннотация", "abstract", "в работе", "in the paper"]):
            abstract_text = p
            continue
        if len(p) > 400 and not main_body:
            main_body = p

    sem_model = None
    try:
        if settings.check_translation or settings.check_abstract:
            print("Загрузка семантической модели (rubert-tiny2)...", flush=True)
            log_memory("Перед загрузкой rubert-tiny2")
            
            # ВАЖНО: Используем ультра-легкую модель!
            sem_model = SentenceTransformer('cointegrated/rubert-tiny2', device="cpu")
            
            log_memory("После загрузки rubert-tiny2")

            if settings.check_translation and org_ru and org_en:
                emb1 = sem_model.encode(org_ru, convert_to_tensor=True)
                emb2 = sem_model.encode(org_en, convert_to_tensor=True)
                sim = util.pytorch_cos_sim(emb1, emb2).item()
                # Порог снижен до 0.6, так как у tiny модели другое распределение векторов
                if sim < 0.60:
                    errors.append(f"[NLP] Перевод организации: сходство {int(sim*100)}%")
            
            if settings.check_abstract and abstract_text and main_body:
                emb_a = sem_model.encode(abstract_text, convert_to_tensor=True)
                emb_m = sem_model.encode(main_body[:1000], convert_to_tensor=True)
                sim = util.pytorch_cos_sim(emb_a, emb_m).item()
                if sim < 0.35:
                    errors.append(f"[NLP] Аннотация: слабое соответствие ({int(sim*100)}%)")
                
    except Exception as e:
        print(f"Ошибка в блоке семантики: {e}", flush=True)
        errors.append(f"[Система] Ошибка NLP модуля")
    finally:
        if sem_model is not None:
            del sem_model
        gc.collect()
        print("Семантическая модель выгружена.", flush=True)

    if settings.check_expert and not any(word in full_text_lower for word in["экспертное заключение", "экспортный контроль"]):
        errors.append("[Экспертиза] Не найдено упоминание об экспертном заключении")

    return errors

# Зависимость для получения сессии БД
def get_db():
    db = database.SessionLocal()
    try:
        yield db
    finally:
        db.close()

@app.post("/register")
def register(email: str, password: str, db: Session = Depends(get_db)):
    # Bcrypt не принимает пароли длиннее 72 символов
    if len(password.encode('utf-8')) > 72:
        return {"error": "Пароль слишком длинный (максимум 72 символа)"}
        
    if db.query(database.User).filter(database.User.email == email).first():
        return {"error": "Email уже занят"}
    
    new_user = database.User(email=email, hashed_password=get_password_hash(password))
    db.add(new_user)
    db.commit()
    return {"message": "Успешно"}

# @app.post("/login")
# def login(email: str, password: str, response: Response, db: Session = Depends(get_db)):
#     user = db.query(database.User).filter(database.User.email == email).first()
#     if not user or not verify_password(password, user.hashed_password):
#         return {"error": "Неверные данные"}
    
#     token = jwt.encode({"sub": user.email}, SECRET_KEY, algorithm=ALGORITHM)
    
#     # Устанавливаем куку
#     response.set_cookie(
#         key="access_token", 
#         value=token, 
#         httponly=True, 
#         max_age=604800, # 7 дней
#         samesite="none" if IS_PRODUCTION else "lax",
#         secure=True if IS_PRODUCTION else False,
#         domain=None
#     )
#     return {"email": user.email, "role": user.role }
@app.post("/login")
def login(email: str, password: str, db: Session = Depends(get_db)):
    user = db.query(database.User).filter(database.User.email == email).first()
    if not user or not verify_password(password, user.hashed_password):
        return {"error": "Неверные данные"}
    
    token = jwt.encode({"sub": user.email}, SECRET_KEY, algorithm=ALGORITHM)
    
    # Больше не устанавливаем куку!
    return {"email": user.email, "role": user.role, "access_token": token}

@app.post("/logout")
def logout(response: Response):
    response.delete_cookie(
        "access_token",
        samesite="none" if IS_PRODUCTION else "lax",
        secure=True if IS_PRODUCTION else False,
        domain=None
    )
    return {"message": "Вышли"}

@app.get("/me")
def me(user: database.User = Depends(get_current_user)):
    if not user: return {"error": "Not logged in"}
    return {"email": user.email, "role": user.role}

@app.get("/")
def home():
    return {"message": "AI Detector API is running"}

@app.post("/feedback/{result_id}")
def save_feedback(result_id: int, correct: bool, db: Session = Depends(get_db), user: database.User = Depends(get_current_user)):
    if not user: raise HTTPException(status_code=401)
    
    db_result = db.query(database.DetectionResult).filter(
        database.DetectionResult.id == result_id,
        database.DetectionResult.owner_id == user.id
    ).first()
    
    if not db_result or db_result.user_feedback is not None:
        raise HTTPException(status_code=404, detail="Результат не найден или фидбек уже учтен")

    # Определяем метку (1 - AI, 0 - Human)
    if correct:
        final_label = 1 if db_result.label == "AI" else 0
    else:
        final_label = 0 if db_result.label == "AI" else 1

    # Запись в новую таблицу
    new_train_data = database.TrainingData(
        text_content=db_result.text_content, 
        label=final_label
    )
    db.add(new_train_data)
    
    # Сохраняем статус в базу данных
    db_result.user_feedback = correct
    db.commit()
    
    return {"status": "success"}

@app.get("/stats")
def get_stats(db: Session = Depends(get_db)):
    total = db.query(database.DetectionResult).filter(database.DetectionResult.user_feedback != None).count()
    correct = db.query(database.DetectionResult).filter(database.DetectionResult.user_feedback == True).count()
    
    accuracy = (correct / total * 100) if total > 0 else 0
    return {"user_accuracy": f"{accuracy:.1f}%", "total_feedbacks": total}

@app.post("/analyze")
def analyze_text_endpoint(
    text: str, 
    db: Session = Depends(get_db), 
    user: database.User = Depends(get_current_user)
):
    if not user: 
        raise HTTPException(status_code=401, detail="Необходимо войти в систему")
    
    # 1. Считаем ИИ (теперь получаем 4 значения)
    label, score, chunks = run_ai_logic(text)
    # 2. Сохраняем в базу данных (filename у нас тут None, так как это просто текст)
    db_result = database.DetectionResult(
        text_content=text,
        filename=None, 
        label=label,
        score=float(score),
        owner_id=user.id,
        format_errors=None, # Для простого текста нормоконтроль не проводим
        page_count=0
    )
    
    db.add(db_result)
    db.commit()
    db.refresh(db_result)
    
    # 3. Возвращаем результат фронтенду
    return {
        "id": db_result.id, 
        "label": label, 
        "score": score, 
        "chunks_analyzed": chunks
    }

@app.post("/norm-control")
async def norm_control(file: UploadFile = File(...), db: Session = Depends(get_db), user: database.User = Depends(get_current_user)):
    if not user: raise HTTPException(status_code=401)
    
    # Получаем настройки пользователя (или создаем дефолтные)
    settings = db.query(database.CheckSettings).filter(database.CheckSettings.owner_id == user.id).first()
    if not settings:
        settings = database.CheckSettings(owner_id=user.id)
        db.add(settings)
        db.commit()

    content = await file.read()
    doc = docx.Document(io.BytesIO(content))
    
    errors = []
    
    # 1. Проверка полей (Margins)
    section = doc.sections[0]
    # Переводим из EMU в см (1 см = 360000 EMU)
    if round(section.left_margin.cm, 1) != settings.margin_left:
        errors.append(f"Поле слева: {round(section.left_margin.cm, 1)}см (нужно {settings.margin_left}см)")
    
    # 2. Проверка шрифта и размера (по первому абзацу текста)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip(): # Пропускаем пустые строки
            # Проверка шрифта
            font = paragraph.runs[0].font
            if font.name and font.name != settings.font_name:
                errors.append(f"Шрифт: обнаружен {font.name} (нужен {settings.font_name})")
            
            # Проверка размера
            if font.size and font.size.pt != settings.font_size:
                errors.append(f"Размер шрифта: {font.size.pt}pt (нужно {settings.font_size}pt)")
            break # Проверяем только начало для примера

    # 3. Проверка литературы (Блок 4)
    # Ищем заголовок "Библиографический список" или "Список литературы"
    refs_count = 0
    found_refs = False
    for p in doc.paragraphs:
        if "библиографический" in p.text.lower() or "литература" in p.text.lower():
            found_refs = True
            continue
        if found_refs and p.text.strip():
            refs_count += 1
            
    if refs_count < settings.min_references:
        errors.append(f"Список источников: найдено {refs_count} (минимум {settings.min_references})")

    return {
        "status": "success" if not errors else "warning",
        "errors": errors,
        "checked_params": {
            "font": settings.font_name,
            "size": settings.font_size,
            "margins": f"{settings.margin_left}/{settings.margin_right}"
        }
    }

@app.post("/settings/reset")
def reset_settings(db: Session = Depends(get_db), user: database.User = Depends(get_current_user)):
    if not user: 
        raise HTTPException(status_code=401)
    
    db_settings = db.query(database.CheckSettings).filter(database.CheckSettings.owner_id == user.id).first()
    
    if not db_settings:
        db_settings = database.CheckSettings(owner_id=user.id)
        db.add(db_settings)
    
    # Устанавливаем эталонные значения АПАК
    db_settings.font_name = "Times New Roman"
    db_settings.font_size = 12
    db_settings.margin_left = 2.0
    db_settings.margin_right = 2.0
    db_settings.margin_top = 2.5
    db_settings.margin_bottom = 2.5
    db_settings.min_references = 3
    db_settings.check_translation = True
    db_settings.check_abstract = True
    db_settings.check_expert = False
    db_settings.feedback_enabled = True
    db_settings.check_apak = True
    
    db.commit()
    db.refresh(db_settings)
    return db_settings

@app.get("/settings")
def get_settings(db: Session = Depends(get_db), user: database.User = Depends(get_current_user)):
    if not user: raise HTTPException(status_code=401)
    settings = db.query(database.CheckSettings).filter(database.CheckSettings.owner_id == user.id).first()
    if not settings:
        settings = database.CheckSettings(owner_id=user.id)
        db.add(settings)
        db.commit()
    return settings

@app.post("/settings")
def update_settings(new_settings: dict, db: Session = Depends(get_db), user: database.User = Depends(get_current_user)):
    if not user: 
        raise HTTPException(status_code=401)
    
    # 1. Ищем существующие настройки
    db_settings = db.query(database.CheckSettings).filter(database.CheckSettings.owner_id == user.id).first()
    
    # 2. Если их нет — создаем объект (но еще не сохраняем)
    if not db_settings:
        db_settings = database.CheckSettings(owner_id=user.id)
        db.add(db_settings)

    # 3. Обновляем поля из пришедшего словаря
    for key, value in new_settings.items():
        # Проверяем, есть ли такое поле в модели, и не пытаемся ли мы сменить ID
        if hasattr(db_settings, key) and key not in ['id', 'owner_id']:
            setattr(db_settings, key, value)
            
    try:
        db.commit() # Сохраняем изменения (INSERT или UPDATE произойдет автоматически)
        db.refresh(db_settings)
    except Exception as e:
        db.rollback()
        print(f"Ошибка БД: {e}")
        raise HTTPException(status_code=500, detail="Ошибка при сохранении в базу данных")
        
    return db_settings

def convert_to_pdf_linux(input_path, output_dir):
    """Универсальная конвертация в PDF через LibreOffice для Linux"""
    import subprocess
    import os
    import time
    
    try:
        env = os.environ.copy()
        env['HOME'] = '/tmp'
        
        # Запускаем LibreOffice
        subprocess.run([
            'soffice', '--headless', '--convert-to', 'pdf', 
            '--outdir', output_dir, input_path
        ], check=True, env=env, timeout=30)
        return True
    except Exception as e:
        print(f"Ошибка LibreOffice: {e}", flush=True)
        return False
    finally:
        # ПРИНУДИТЕЛЬНО УБИВАЕМ ПРОЦЕССЫ И ОСВОБОЖДАЕМ ПАМЯТЬ
        print("Убиваем LibreOffice для очистки RAM...", flush=True)
        try:
            # Жестко завершаем все процессы soffice
            subprocess.run(['pkill', '-9', '-f', 'soffice'], stderr=subprocess.DEVNULL)
            subprocess.run(['pkill', '-9', '-f', 'oosplash'], stderr=subprocess.DEVNULL)
            # Ждем 1 секунду, чтобы ОС успела освободить страницы памяти
            time.sleep(1)
        except:
            pass

def convert_doc_to_docx(file_bytes: bytes) -> bytes:
    """Конвертирует старый .doc в современный .docx"""
    if IS_WINDOWS:
        # Логика для Windows (локально)
        with tempfile.TemporaryDirectory() as temp_dir:
            doc_path = os.path.abspath(os.path.join(temp_dir, "temp.doc"))
            docx_path = os.path.abspath(os.path.join(temp_dir, "temp.docx"))
            with open(doc_path, "wb") as f: f.write(file_bytes)
            try:
                pythoncom.CoInitialize()
                word = win32com.client.DispatchEx("Word.Application")
                word.Visible = False
                wb_doc = word.Documents.Open(doc_path)
                wb_doc.SaveAs(docx_path, FileFormat=16) # 16 = wdFormatXMLDocument
                wb_doc.Close(0)
                word.Quit()
                if os.path.exists(docx_path):
                    with open(docx_path, "rb") as f: return f.read()
            except Exception as e:
                print(f"Windows DOC error: {e}", flush=True)
    else:
        # Логика для Linux (Railway) через LibreOffice
        with tempfile.TemporaryDirectory() as temp_dir:
            in_path = os.path.join(temp_dir, "temp.doc")
            with open(in_path, "wb") as f: f.write(file_bytes)
            try:
                # В Linux конвертируем .doc -> .docx напрямую через soffice
                subprocess.run(['soffice', '--headless', '--convert-to', 'docx', '--outdir', temp_dir, in_path], check=True, timeout=30)
                out_path = os.path.join(temp_dir, "temp.docx")
                if os.path.exists(out_path):
                    with open(out_path, "rb") as f: return f.read()
            except Exception as e:
                print(f"Linux DOC error: {e}", flush=True)
    return file_bytes

def get_page_count(file_bytes: bytes, filename: str) -> int:
    """Считает количество страниц в документе"""
    ext = filename.split(".")[-1].lower()
    try:
        if ext == "pdf":
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            return doc.page_count
            
        with tempfile.TemporaryDirectory() as temp_dir:
            file_path = os.path.join(temp_dir, f"temp.{ext}")
            pdf_path = os.path.join(temp_dir, "temp.pdf")
            with open(file_path, "wb") as f: f.write(file_bytes)

            if IS_WINDOWS:
                if ext == "docx":
                    # Используем быструю docx2pdf
                    convert(file_path, pdf_path)
                elif ext == "rtf":
                    pythoncom.CoInitialize()
                    word = win32com.client.DispatchEx("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(os.path.abspath(file_path))
                    doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17) # 17 = PDF
                    doc.Close(0)
                    word.Quit()
                elif ext == "txt":
                    return max(1, len(file_bytes) // 3000)
            else:
                # В облаке (Linux) всё делаем через одну команду LibreOffice
                convert_to_pdf_linux(file_path, temp_dir)
            
            if os.path.exists(pdf_path):
                temp_pdf = fitz.open(pdf_path)
                pages = temp_pdf.page_count
                temp_pdf.close()
                return pages
        
        # Если ничего не помогло - считаем по символам
        return max(1, len(file_bytes) // 2500)
                
    except Exception as e:
        print(f"Ошибка подсчета страниц: {e}", flush=True)
        return 1

def process_docx_apak(file_bytes: bytes, settings: database.CheckSettings):
    doc = docx.Document(io.BytesIO(file_bytes))
    errors =[]
    html_lines =[]
    
    current_block = "header" 
    empty_lines = 0
    valid_single_letters =['а', 'и', 'в', 'о', 'у', 'с', 'к', 'я', 'б', 'ж', 'z', 'a', 'i']
    
    # Флаг: вошли ли мы в секцию библиографии
    in_references_section = False

    for p in doc.paragraphs:
        raw_text = p.text.replace('\t', '    ').replace('\x0c', '')
        stripped_text = raw_text.strip()
        
        has_image = len(p._element.xpath('.//w:drawing')) > 0 or len(p._element.xpath('.//w:pict')) > 0

        if not stripped_text:
            if not has_image:
                empty_lines += 1
            html_lines.append("<div style='height: 1em;'></div>")
            continue
            
        lower_text = stripped_text.lower()
        
        # --- 1. УМНОЕ ОПРЕДЕЛЕНИЕ БЛОКА И ВЫРАВНИВАНИЯ ---
        alignment = "justify" 
        indent = "1.25cm" # Стандартный абзацный отступ ГОСТ
        is_bold_override = False
        is_spacing_error = False 
        is_figure_caption = stripped_text.startswith("Рис.") or stripped_text.startswith("Fig.")
        is_list_item = bool(p._element.xpath('.//w:numPr'))
        
        # 1. УДК (Слева)
        if lower_text.startswith("удк") or lower_text.startswith("udc"):
            current_block = "udk"
            alignment = "left"
            indent = "0"
            
        # 2. Копирайт (Справа)
        elif "©" in lower_text or "(c)" in lower_text:
            current_block = "copyright"
            alignment = "right"
            indent = "0"
            if settings.norm_enabled and settings.check_apak and empty_lines != 1:
                errors.append(f"[АПАК] Перед копирайтом нужна 1 пустая строка (найдено: {empty_lines})")
                is_spacing_error = True

        # 3. Заголовки "Библиографические ссылки" / "Список литературы" (По центру, жирным)
        elif len(stripped_text) < 80 and (("библиографическ" in lower_text and "ссылк" in lower_text) or "список литературы" in lower_text or "references" in lower_text):
            current_block = "references_header"
            alignment = "center"
            indent = "0"
            in_references_section = True
            is_bold_override = True # Принудительно делаем жирным
            
        # 4. Если мы вошли в секцию ссылок, то сами ссылки по ширине без отступа
        elif in_references_section:
            current_block = "reference_item"
            alignment = "justify"
            indent = "0" 
            
        # 5. Аннотация и ключевые слова (По центру)
        elif lower_text.startswith("аннотация") or lower_text.startswith("abstract") or lower_text.startswith("ключевые слова") or lower_text.startswith("keywords"):
            current_block = "abstract"
            alignment = "center"
            indent = "0"
            
        # 6. Картинки и подписи (По центру)
        elif is_figure_caption or has_image:
            current_block = "figure"
            alignment = "center"
            indent = "0"
            
        # 7. Университет / Email (По центру)
        elif any(x in lower_text for x in["сибирский государственный", "reshetnev", "university", "федеральное", "г. красноярск", "krasnoyarsk", "просп.", "prospekt", "e-mail", "mail.ru", "yandex.ru", "gmail.com"]):
            current_block = "university"
            alignment = "center"
            indent = "0"
            
        # 8. ФИО, Руководитель, Название статьи, Заголовки (Всё короткое -> По центру)
        elif len(stripped_text) < 100 and not is_list_item and not stripped_text.endswith((";", ":", ",", "»", '"', "-")):
            current_block = "heading_or_title"
            alignment = "center"
            indent = "0"
            
        # 9. Основной длинный текст (По ширине - Justify)
        else:
            current_block = "main"
            alignment = "justify"
            indent = "1.25cm"

        # --- 2. УСТАНОВКА ОЖИДАЕМОГО РАЗМЕРА ---
        expected_size = 11 if current_block == "university" else settings.font_size
        empty_lines = 0

        # --- 3. СТИЛЬ АБЗАЦА ---
        p_style = f"margin: 0; line-height: 1.5; white-space: pre-wrap; vertical-align: baseline; font-family: 'Times New Roman', serif; "
        p_style += f"text-align: {alignment}; text-indent: {indent}; "
        
        if is_bold_override:
            p_style += "font-weight: bold; "
            
        if is_spacing_error:
            p_style += "background-color: #fef08a; outline: 2px dashed #ca8a04; border-radius: 2px; "

        p_html = f'<p style="{p_style}">'
        
        # ВАЖНО: Эта переменная должна быть ровно здесь, ПЕРЕД циклом runs!
        marker_added = False 
        
        for run in p.runs:
            if not run.text: continue
            t = run.text.replace("<", "&lt;").replace(">", "&gt;")
            
            # Восстанавливаем маркер списка (тире)
            if is_list_item and not marker_added and t.strip():
                if not t.strip().startswith(("–", "-", "—", "•")):
                    t = f"–&nbsp;&nbsp;&nbsp;{t}"
                marker_added = True
                
            f_size = None
            has_explicit_size = False
            
            if run.font and run.font.size:
                f_size = run.font.size.pt
                has_explicit_size = True
            elif p.style and hasattr(p.style, 'font') and p.style.font and p.style.font.size:
                f_size = p.style.font.size.pt
                if p.style.name != 'Normal':
                    has_explicit_size = True
            elif p.style and hasattr(p.style, 'base_style') and p.style.base_style and hasattr(p.style.base_style, 'font') and p.style.base_style.font and p.style.base_style.font.size:
                f_size = p.style.base_style.font.size.pt
                if p.style.base_style.name != 'Normal':
                    has_explicit_size = True
            else:
                try: 
                    if doc.styles['Normal'].font.size:
                        f_size = doc.styles['Normal'].font.size.pt
                except: pass
            
            if f_size is not None and round(f_size) == 11 and not has_explicit_size:
                f_size = expected_size

            # Проверка шрифта (только если нормоконтроль Включен)
            if settings.norm_enabled and current_block != "figure" and f_size and abs(round(f_size) - expected_size) > 0.1:
                errors.append(f"[Шрифт] Ожидался {expected_size}pt, найден {round(f_size)}pt")
                t = f"<mark style='background-color: #fecaca; color: #991b1b; padding: 0;' title='Ожидался {expected_size}pt, найден {round(f_size)}pt'>{t}</mark>"

            # Проверка пробелов (только если нормоконтроль Включен)
            if settings.norm_enabled and settings.check_apak:
                if "  " in t:
                    errors.append(f"[Пробелы] Лишние пробелы")
                    t = t.replace("  ", "<span style='background-color: #fef08a; box-shadow: 0 2px 0 #ca8a04;' title='Лишние пробелы'>  </span>")

                def sub_standalone(m):
                    letter = m.group(2)
                    if letter.lower() not in valid_single_letters:
                        errors.append(f"[Типографика] Разрыв слова '{letter}'")
                        return f"{m.group(1)}<mark style='background-color: #ffedd5; outline: 1px solid #ea580c; outline-offset: -1px;'>{letter}</mark> "
                    return m.group(0)
                t = re.sub(r'(^|\s)([а-яА-ЯёЁ])\s', sub_standalone, t)

                def sub_illegal_start(m):
                    errors.append(f"[Типографика] Разрыв слова")
                    return f"{m.group(1)}<mark style='background-color: #ffedd5; outline: 1px solid #ea580c;'>{m.group(2)}</mark>{m.group(3)}"
                t = re.sub(r'(^|\s)([ьыъЬЫЪ])([а-яА-ЯёЁ]+)', sub_illegal_start, t)

            if run.bold: t = f"<b>{t}</b>"
            if run.italic: t = f"<i>{t}</i>"
            p_html += t
            
        p_html += "</p>"
        html_lines.append(p_html)

    return "".join(html_lines), list(dict.fromkeys(errors))

# эндпоинт для загрузки файлов
@app.post("/analyze-file")
async def analyze_file(file: UploadFile = File(...), db: Session = Depends(get_db), user: database.User = Depends(get_current_user)):
    if not user: raise HTTPException(status_code=401)
    
    # 1. Получаем настройки и читаем файл
    settings = db.query(database.CheckSettings).filter(database.CheckSettings.owner_id == user.id).first()
    file_bytes = await file.read()
    filename = file.filename
    ext = filename.split(".")[-1].lower()
    
    # 2. Если файл .doc — незаметно превращаем его в .docx для работы парсера
    log_memory("До конвертации DOC")
    if ext == "doc":
        converted_bytes = convert_doc_to_docx(file_bytes)
        if converted_bytes != file_bytes:
            file_bytes = converted_bytes
            ext = "docx"
            filename = filename + "x"
    log_memory("После конвертации DOC")
    # 3. Извлекаем чистый текст для ИИ-детектора
    text = await extract_text_from_file_bytes(file_bytes, filename)
    
    # 4. Инициализируем переменные результата
    html_content = ""
    format_errors = []
    label, score, chunks = "Disabled", 0.0, 0
    
    # 5. ЗАПУСКАЕМ УМНЫЙ ПАРСЕР АПАК (только для DOCX/DOC)
    if ext == "docx":
        # Наш новый парсер возвращает готовый HTML с подсветкой и список технических ошибок
        html_content, format_errors = process_docx_apak(file_bytes, settings)
        
        # Если включен нормоконтроль — добавляем семантические проверки (перевод, аннотация)
        if settings.norm_enabled:
            doc_obj = docx.Document(io.BytesIO(file_bytes))
            semantic_errors = check_semantic_rules(doc_obj, settings)
            format_errors.extend(semantic_errors)
    else:
        # Для PDF или TXT просто выводим текст в теге <pre>, чтобы сохранить пробелы
        html_content = f"<div style='white-space: pre-wrap; font-family: serif;'>{text}</div>"
    
    import gc
    gc.collect() 
    
    # 6. Проверка на ИИ (Детектор)
    if settings.ai_enabled:
        # run_ai_logic теперь возвращает 3 параметра (без подсветки ИИ, так как мы её убрали)
        label, score, chunks = run_ai_logic(text)
        
    # 7. Считаем страницы
    page_count = get_page_count(file_bytes, filename)

    # 8. Сохраняем итоговый результат в базу данных
    db_result = database.DetectionResult(
        text_content=text,
        html_content=html_content, # Сохраняем HTML с красной подсветкой АПАК
        filename=filename,
        label=label,
        score=float(score),
        format_errors=format_errors,
        page_count=page_count,
        owner_id=user.id
    )
    db.add(db_result)
    db.commit()
    db.refresh(db_result)
    
    # 9. Отдаем фронтенду
    return {
        "id": db_result.id, 
        "label": label, 
        "score": score, 
        "format_errors": format_errors, 
        "html_content": html_content, 
        "page_count": page_count, 
        "filename": filename
    }

@app.get("/history")
def get_history(db: Session = Depends(get_db), user: database.User = Depends(get_current_user)):
    if not user: return []
    return db.query(database.DetectionResult).filter(database.DetectionResult.owner_id == user.id).all()

# ==========================================
# ФУНКЦИИ УДАЛЕНИЯ ИСТОРИИ И АДМИН-ПАНЕЛЬ
# ==========================================

# СНАЧАЛА СТАТИЧЕСКИЙ ПУТЬ
@app.delete("/history/all")
def delete_all_my_history(db: Session = Depends(get_db), user: database.User = Depends(get_current_user)):
    if not user: raise HTTPException(status_code=401)
    db.query(database.DetectionResult).filter(database.DetectionResult.owner_id == user.id).delete()
    db.commit()
    return {"message": "Вся ваша история очищена"}

# 2. ПОТОМ ПУТЬ С ПАРАМЕТРОМ
@app.delete("/history/{result_id}")
def delete_history_item(result_id: int, db: Session = Depends(get_db), user: database.User = Depends(get_current_user)):
    if not user: raise HTTPException(status_code=401)
    item = db.query(database.DetectionResult).filter(
        database.DetectionResult.id == result_id, 
        database.DetectionResult.owner_id == user.id
    ).first()
    if not item: raise HTTPException(status_code=404, detail="Скан не найден")
    db.delete(item)
    db.commit()
    return {"message": "Удалено"}

# --- ЗАЩИТА АДМИНКИ ---
def get_admin_user(user: database.User = Depends(get_current_user)):
    if not user:
        raise HTTPException(status_code=401, detail="Не авторизован")
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Доступ запрещен. Вы не администратор.")
    return user

import pandas as pd
from fastapi.responses import StreamingResponse
import io

@app.get("/admin/export-dataset")
def export_dataset(admin: database.User = Depends(get_admin_user), db: Session = Depends(get_db)):
    # 1. Забираем все данные из таблицы training_data
    data = db.query(database.TrainingData).all()
    if not data:
        raise HTTPException(status_code=404, detail="Датасет пуст")
        
    # 2. Формируем файл
    df = pd.DataFrame([{"text": d.text_content, "label": d.label} for d in data])
    stream = io.StringIO()
    # Используем quotechar, чтобы переносы строк внутри текста не ломали структуру CSV
    df.to_csv(stream, index=False, encoding='utf-8-sig', quoting=csv.QUOTE_ALL)
    
    # 3. ОЧИСТКА БАЗЫ (выполняем после успешного формирования)
    db.query(database.TrainingData).delete()
    db.commit()
    
    response = Response(content=stream.getvalue(), media_type="text/csv")
    response.headers["Content-Disposition"] = "attachment; filename=training_dataset.csv"
    return response

# Получить всех пользователей (для админки)
@app.get("/admin/users")
def admin_get_users(db: Session = Depends(get_db), admin: database.User = Depends(get_admin_user)):
    users = db.query(database.User).all()
    # Возвращаем юзеров и количество их сканов
    return[{"id": u.id, "email": u.email, "role": u.role, "scans_count": len(u.results)} for u in users]

# Создать нового админа или пользователя
@app.post("/admin/users")
def admin_create_user(email: str, password: str, role: str = "user", db: Session = Depends(get_db), admin: database.User = Depends(get_admin_user)):
    if db.query(database.User).filter(database.User.email == email).first():
        raise HTTPException(status_code=400, detail="Email занят")
    new_user = database.User(email=email, hashed_password=get_password_hash(password), role=role)
    db.add(new_user)
    db.commit()
    return {"message": f"Пользователь {email} ({role}) создан!"}

# Удалить любого пользователя (для админки)
@app.delete("/admin/users/{target_id}")
def admin_delete_user(target_id: int, db: Session = Depends(get_db), admin: database.User = Depends(get_admin_user)):
    user_to_delete = db.query(database.User).filter(database.User.id == target_id).first()
    if not user_to_delete: raise HTTPException(status_code=404)
    # Сначала удаляем все его сканы, настройки и отзывы
    db.query(database.DetectionResult).filter(database.DetectionResult.owner_id == target_id).delete()
    db.query(database.CheckSettings).filter(database.CheckSettings.owner_id == target_id).delete()
    db.delete(user_to_delete)
    db.commit()
    return {"message": "Пользователь и все его данные удалены"}

# Удалить ВСЮ историю ВООБЩЕ У ВСЕХ (кнопка паники)
@app.delete("/admin/history/wipe-all")
def admin_wipe_all_history(db: Session = Depends(get_db), admin: database.User = Depends(get_admin_user)):
    db.query(database.DetectionResult).delete()
    db.commit()
    return {"message": "Глобальная база сканов полностью очищена"}

@app.delete("/admin/users/{target_id}/history")
def admin_clear_user_history(target_id: int, db: Session = Depends(get_db), admin: database.User = Depends(get_admin_user)):
    db.query(database.DetectionResult).filter(database.DetectionResult.owner_id == target_id).delete()
    db.commit()
    return {"message": f"История пользователя #{target_id} очищена"}